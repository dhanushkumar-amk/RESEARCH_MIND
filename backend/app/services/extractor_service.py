import io
import time
import pdfplumber
import PyPDF2
from docx import Document
import openpyxl
from fastapi import HTTPException
import httpx

class ExtractorService:
    def extract_text(self, file_bytes: bytes, file_type: str) -> list[dict]:
        """
        Main entry point for document extraction.
        Routes to the correct extractor based on file type.
        Returns a list of dicts: [{"text": str, "page_number": int}]
        """
        file_type = file_type.lower().strip(".")
        
        if file_type == "pdf":
            return self.extract_pdf(file_bytes)
        elif file_type == "docx":
            return [{"text": self.extract_docx(file_bytes), "page_number": 1}]
        elif file_type == "xlsx":
            return [{"text": self.extract_xlsx(file_bytes), "page_number": 1}]
        elif file_type in {"txt", "text"}:
            return [{"text": self.extract_txt(file_bytes), "page_number": 1}]
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type for text extraction: {file_type}"
            )

    def extract_pdf(self, file_bytes: bytes) -> list[dict]:
        """
        Step 2: PDF extraction using pdfplumber with PyPDF2 as a fallback.
        Preserves individual page numbers.
        """
        pages = []
        # Try pdfplumber first (better formatting)
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages.append({
                            "text": page_text.strip(),
                            "page_number": i + 1
                        })
        except Exception as e:
            print(f"pdfplumber failed: {e}. Trying fallback PyPDF2...")
            pages = []
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages.append({
                            "text": page_text.strip(),
                            "page_number": i + 1
                        })
            except Exception as fe:
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to extract text from PDF: {str(fe)}"
                )

        if not pages:
            raise HTTPException(
                status_code=400,
                detail="PDF appears to be empty or contains only scanned images (no selectable text)."
            )
        return pages

    def extract_docx(self, file_bytes: bytes) -> str:
        """
        Step 3: Microsoft Word (.docx) extraction.
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text.append(paragraph.text)
            
            # Extract from tables as well
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text.append(" | ".join(row_text))
                        
            clean_text = "\n".join(text).strip()
            if not clean_text:
                raise HTTPException(status_code=400, detail="Word document is empty.")
            return clean_text
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from Word document: {str(e)}"
            )

    def extract_xlsx(self, file_bytes: bytes) -> str:
        """
        Step 3: Microsoft Excel (.xlsx) extraction.
        """
        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            text_lines = []
            
            for sheet in wb.worksheets:
                text_lines.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    row_values = [str(cell).strip() if cell is not None else "" for cell in row]
                    if any(row_values):
                        text_lines.append(" | ".join(row_values))
                        
            clean_text = "\n".join(text_lines).strip()
            if not clean_text:
                raise HTTPException(status_code=400, detail="Excel sheet is empty.")
            return clean_text
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from Excel document: {str(e)}"
            )

    def extract_txt(self, file_bytes: bytes) -> str:
        """
        Step 3: Plain Text (.txt) extraction with encoding detection.
        """
        encodings = ["utf-8", "latin-1", "utf-16", "cp1252"]
        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                clean_text = text.strip()
                if not clean_text:
                    raise HTTPException(status_code=400, detail="Text file is empty.")
                return clean_text
            except UnicodeDecodeError:
                continue
        
        raise HTTPException(
            status_code=400,
            detail="Failed to decode text file. Ensure it is encoded in UTF-8 or standard ASCII."
        )

    def _get_with_retry(self, url: str, headers: dict = None, retries: int = 3, backoff: int = 2) -> httpx.Response:
        """
        Helper method to execute HTTP requests with exponential backoff retries.
        """
        sleep_time = backoff
        for attempt in range(retries):
            try:
                with httpx.Client(timeout=10.0, follow_redirects=True) as client:
                    response = client.get(url, headers=headers)
                    response.raise_for_status()
                    return response
            except (httpx.HTTPStatusError, httpx.RequestError) as e:
                if attempt == retries - 1:
                    raise e
            print(f"HTTP request failed for {url}. Retrying in {sleep_time}s... (Attempt {attempt+1}/{retries})")
            time.sleep(sleep_time)
            sleep_time *= 2

    def extract_url(self, url: str) -> str:
        """
        Step 4: Web HTML scraping using trafilatura with BeautifulSoup as a fallback.
        Includes retry logic for production stability.
        """
        import trafilatura
        from bs4 import BeautifulSoup

        # Try trafilatura first (advanced content extraction, strips ads/navigation)
        try:
            downloaded = trafilatura.fetch_url(url)
            if downloaded:
                result = trafilatura.extract(downloaded)
                if result:
                    return result.strip()
        except Exception as e:
            print(f"trafilatura scraping failed for {url}: {e}. Trying BeautifulSoup fallback...")

        # Fallback with retry logic
        try:
            headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
            response = self._get_with_retry(url, headers=headers)
                
            soup = BeautifulSoup(response.text, "html.parser")
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()
                
            text = soup.get_text(separator="\n")
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            clean_text = "\n".join(chunk for chunk in chunks if chunk)
            
            if not clean_text:
                raise HTTPException(status_code=400, detail="Webpage scraped text is empty.")
            return clean_text
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to scrape webpage text: {str(e)}"
            )

    def extract_youtube(self, url: str) -> str:
        """
        Step 5: YouTube transcript extraction using youtube-transcript-api.
        """
        from youtube_transcript_api import YouTubeTranscriptApi
        from urllib.parse import urlparse, parse_qs
        import re

        # Extract video ID from URL
        video_id = None
        parsed = urlparse(url)
        if "youtube.com" in parsed.netloc:
            if parsed.path == "/watch":
                video_id = parse_qs(parsed.query).get("v", [None])[0]
            elif parsed.path.startswith(("/embed/", "/v/")):
                video_id = parsed.path.split("/")[2]
        elif "youtu.be" in parsed.netloc:
            video_id = parsed.path.lstrip("/")

        # Regex fallback just in case
        if not video_id:
            match = re.search(r"(?:v=|\/)([0-9A-Za-z_-]{11}).*", url)
            if match:
                video_id = match.group(1)

        if not video_id:
            raise HTTPException(
                status_code=400,
                detail="Could not extract a valid YouTube video ID from the provided link."
            )

        try:
            # Try to fetch transcript in English, then fallback to other languages / auto-generated
            if hasattr(YouTubeTranscriptApi, "list"):
                try:
                    transcript_list = YouTubeTranscriptApi().list(video_id)
                except TypeError:
                    transcript_list = YouTubeTranscriptApi.list(video_id)
            else:
                transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
                
            try:
                transcript = transcript_list.find_transcript(['en'])
            except Exception:
                # Fallback to any available language
                transcript = next(iter(transcript_list))
                
            data = transcript.fetch()
            text_pieces = [entry["text"] for entry in data]
            clean_text = " ".join(text_pieces).strip()
            
            if not clean_text:
                raise HTTPException(status_code=400, detail="YouTube video transcript is empty.")
            return clean_text
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Failed to fetch YouTube transcript. Ensure subtitles are enabled for this video. Details: {str(e)}"
            )
